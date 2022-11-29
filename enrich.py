import numpy as np
import pinyin
import requests
from cihai.core import Cihai
import re
import json
from tqdm import tqdm
import os
from cedict_utils.cedict import CedictParser
import zhon
import zhon.cedict
import docx
from docx import Document
from docx.shared import RGBColor, Mm, Pt
from docx.enum.text import WD_BREAK
import typer
import pynlpir
from docx2pdf import convert
#import jieba
#import epitran
from PETRUS import use_petrus

language = "br"

#epi = epitran.Epitran('por-Latn')
#epi.transliterate("bom dia")

# phonetics mapping thanks to http://www.nilc.icmc.usp.br/aeiouado/
# read csv and create dictionary
br_dict = {}
with open("phonetics_mappings/br.csv", "r", encoding="utf8") as f:
    lines = f.readlines()
    lines = [x.strip() for x in lines]
    lines = [x.split("\t") for x in lines]

    for line in lines:
        br_dict[line[0]] = line[1]


phonetic_fs = {  "br": lambda x: use_petrus.get_phonetics(x) #br_dict.get(x.strip().upper(), use_petrus.get_phonetics(x.strip()))
               , "zh": lambda x: pinyin.get(x)}

links_chinese = {"hyperlink_base": "https://www.mdbg.net/chinese/dictionary?page=worddict&wdrst=0&wdqb="
                 , "google_translate": "https://translate.googleapis.com/translate_a/single?client=gtx&sl=zh-CN&tl=en&dt=t&q="
                 }
links_portuguese = {"hyperlink_base": "https://www.collinsdictionary.com/dictionary/english-portuguese/"
                    , "google_translate": "https://translate.googleapis.com/translate_a/single?client=gtx&sl=pt&tl=en&dt=t&q="
                    }

links_all = {"zh": links_chinese, "br": links_portuguese}
links=links_all[language]

class Translator:

    def __init__(self, language):
        self.language = language

    def get_phonetics(self, text):
        ret = phonetic_fs[self.language](text)
        # remove brackets
        ret = re.sub(r"\[.*?\]", "", ret)

        return Character(ret, pinyin_size)

    def translation(self, x):
        # translate from chinese to english using a free translation service
        url = links["google_translate"] + x
        response = requests.get(url)
        if response.status_code != 200:
            raise Exception("ERROR: API request unsuccessful.")
        ret = json.loads(response.content.decode("utf-8"))[0][0][0]
        return Character(ret, translation_size)


    def tokenize(self, line, engine):

        if self.language == "zh":
            if engine == "jieba":

                words = jieba.lcut(line)
                tokenized = list(zip(words, len(words) * ["not_detected"]))
                return tokenized

            elif engine == "pynlpir":
                pynlpir.open()
                tokenized = pynlpir.segment(line)
                pynlpir.close()
                return tokenized

        elif self.language == "br":
            words = [" " + x for x in line.split(" ")]
            tokenized = list(zip(words, len(words) * ["not_detected"]))
            return tokenized

        else:
            raise Exception("ERROR: Language not supported.")

translator = Translator(language)


debug = False
base_size = 40
pinyin_size = base_size / 2
word_for_word_size = base_size / 4
translation_size = base_size / 2
font_name = "Courier New"
chin = "MS Mincho"
ratio_lucida_to_mincho = 50 / 30
space_size = 1
lines_per_page = 14

color_coding = {
    "noun": RGBColor(17, 138, 178),  # blue
    "verb": RGBColor(219, 58, 52),  # red
    "adjective": RGBColor(6, 214, 160),  # green
    "adverb": RGBColor(247, 127, 0),  # yellow
    "classifier": RGBColor(0, 48, 73),  # dark blue
    "pronoun": RGBColor(239, 71, 111),  # pink
}

replace_dict = {
    "“": '"',
    "”": '"',
    "‘": "'",
    "’": "'",
    "—": "-",
    "…": "...",
    "《": "<",
    "》": ">",
    "（": "(",
    "）": ")",
    "【": "[",
    "】": "]",
    "、": ",",
    "。": ".",
    "，": ",",
    "：": ":",
    "；": ";",
    "？": "?",
    "！": "!",
    "\n": " "
}


# load dictionary
parser = CedictParser()
entries = parser.parse()
dictionary = {}
for e in entries[::-1]:
    dictionary[e.simplified] = e.meanings[0]

c = Cihai()


class Character:

    def __init__(self, text, size):
        self.text = text
        self.size = size
        # check if character is chinese
        punctuation = re.findall('[{}]'.format(zhon.hanzi.punctuation), text)
        simplified = re.findall('[{}]'.format(zhon.cedict.simplified), text)
        self.chinese = len(punctuation + simplified) > 0
        self.color = RGBColor(0, 0, 0)  # black

    def len(self):
        len_ = len(self.text) * self.size
        if not self.chinese:
            return len_
        else:
            return len_ * ratio_lucida_to_mincho

    def set_color(self, color):
        self.color = color


def word_for_word(x, notranslate=False):
    if x in notranslate:
        return Character("", word_for_word_size)

    try:
        if x not in dictionary:
            trans = translator.translation(x).text
        elif dictionary[x][:10] == "variant of":
            trans = translator.translation(x).text
        elif dictionary[x][:7] == "surname":
            trans = translator.translation(x).text
        else:
            trans = dictionary[x]

        # remove everything between parentheses
        trans = re.sub(r"\(.*\)", "", trans)

        if trans[:6] == "to be ":
            trans = trans[6:]
        elif trans[:3] == "to ":
            trans = trans[3:]

        splits = trans.split(" ")
        old_l = splits[0]
        for word in splits[1:]:

            new_l = old_l + " " + word
            if len(new_l) > 10:
                break

            old_l = new_l

        ret = old_l[:10]

        if not ret:
            ret = ""

    except:
        ret = ""

    return Character(ret, word_for_word_size)



def pad(x, length):
    if x.len() == length:
        return "", ""
    elif x.len() < length:
        # define left and right pad, round up left, round down right
        left = int(np.ceil((length - x.len()) / 2) / space_size)
        right = int(np.floor((length - x.len()) / 2) / space_size)

        return left * " ", right * " "


document = Document()
section = document.sections[0]
section.page_height = Mm(297 * 3)
section.page_width = Mm(210 * 3)


def preprocess(line):
    for key, value in replace_dict.items():
        line = line.replace(key, value)
    return line + "."


def add_hyperlink(paragraph, text, url, color, font_name, size):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = color
    r.font.name = font_name
    r.font.size = size

    return hyperlink



def enrich_txt(input_path:str, output_path:str, use_notranslate_file:bool=True, engine="pynlpir") -> None:
    """
    Enriches a txt file with pinyin, translation and word for word translation.
    :param input_path: path to input txt file
    :param output_path: path to output docx file
    :param use_notranslate_file: if True, uses notranslate.txt to not translate certain words
    :param engine: engine to use for word tokenization, either "pynlpir" or "jieba"
    :return: None

    """

    # load text file from data
    with open(input_path, "r", encoding="UTF-8") as f:
        data = f.read()
        data = data.replace("。”", "”。")
        data = data.split("。")


    # load list of words that should not be translated
    if use_notranslate_file:
        with open("notranslate.txt", "r", encoding="UTF-8") as f:
            notranslate = f.read().splitlines()
            notranslate = set(notranslate)
    else:
        notranslate = set()


    if not c.unihan.is_bootstrapped:  # download and install Unihan to db
        c.unihan.bootstrap()

    def display_padded_text(text_list, padding, add_dict_link=False):
        for i in range(len(text_list)):

            pad_left, pad_right = padding[i]
            run = p.add_run(pad_left)
            run.font.size = Pt(space_size)
            run.font.name = font_name
            # add hyperlink
            if add_dict_link:
                url = links["hyperlink_base"] + text_list[i].text
                add_hyperlink(p, text_list[i].text, url, text_list[i].color, font_name, Pt(text_list[i].size))
            else:
                run = p.add_run(text_list[i].text)
                run.font.size = Pt(text_list[i].size)
                run.font.name = font_name
                run.font.color.rgb = text_list[i].color

            run = p.add_run(pad_right)
            run.font.size = Pt(space_size)
            run.font.name = font_name

            if debug:
                run = p.add_run("|")
                run.font.size = Pt(15)
                run.font.name = font_name

    line_i = 1
    for line in tqdm(data):
        line = preprocess(line)
        p = document.add_paragraph("")
        run = p.add_run(translator.translation(line).text + "\n")
        run.font.size = Pt(translation_size)
        run.font.name = "Arial"
        wfw_string, pad_wfw = [], []
        py_string, pad_py = [], []
        char_string, pad_char = [], []

        tokenized = translator.tokenize(line, engine)
        sizes = []
        for token, word_type in tokenized:

            wfw = word_for_word(token, notranslate)
            py = translator.get_phonetics(token)
            token = Character(token, base_size)
            token.set_color(color_coding.get(word_type, RGBColor(0, 0, 0)))
            size = max(wfw.len(), py.len(), token.len())
            sizes.append(size)
            pad_wfw.append(pad(wfw, size))
            pad_py.append(pad(py, size))
            pad_char.append(pad(token, size))

            wfw_string.append(wfw)
            py_string.append(py)

            char_string.append(token)

        c_l, c_r = 0, 0
        chunks = []
        while c_l < len(char_string):
            sizes = np.array(sizes)
            cumsum = sizes[c_r:].cumsum()
            # get last value where cumsum is less than 2500

            c_r = c_l + np.where(cumsum < 2500)[0][-1] + 1
            chunk = wfw_string[c_l:c_r], pad_wfw[c_l:c_r], py_string[c_l:c_r], pad_py[c_l:c_r], char_string[c_l:c_r], pad_char[c_l:c_r]
            chunks.append(chunk)
            c_l = c_r

        for wfw_string, pad_wfw, py_string, pad_py, char_string, pad_char in chunks:

            display_padded_text(wfw_string, pad_wfw)
            p.add_run("\n")
            display_padded_text(py_string, pad_py)
            p.add_run("\n")
            display_padded_text(char_string, pad_char, add_dict_link=True)


            if line_i % lines_per_page == 0:
                run = p.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                p.add_run("\n")
            line_i+=1

    document.save(output_path)
    convert(output_path, os.path.abspath(output_path.replace(".docx", ".pdf")))


if __name__ == "__main__":
    #typer.run(enrich_txt)
    enrich_txt("input.txt", "output.docx", use_notranslate_file=True)