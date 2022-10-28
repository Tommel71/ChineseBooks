import numpy as np
import pinyin
import requests
from cihai.core import Cihai
import re
import json
from tqdm import tqdm
from cedict_utils.cedict import CedictParser
import zhon
import zhon.cedict
import docx
from docx import Document
from docx.shared import RGBColor, Mm, Pt
from docx.enum.text import WD_BREAK



hyperlink_base = "https://www.mdbg.net/chinese/dictionary?page=worddict&wdrst=0&wdqb="
debug = False
base_size = 40
pinyin_size = base_size / 2
word_for_word_size = base_size / 4
translation_size = base_size / 2
font_name = "Courier New"
chin = "MS Mincho"
ratio_lucida_to_mincho = 50 / 30
space_size = 1
lines_per_page = 14

color_coding = {
    "noun": RGBColor(17, 138, 178),  # blue
    "verb": RGBColor(219, 58, 52),  # red
    "adjective": RGBColor(6, 214, 160),  # green
    "adverb": RGBColor(247, 127, 0),  # yellow
    "classifier": RGBColor(0, 48, 73),  # dark blue
    "pronoun": RGBColor(239, 71, 111),  # pink
}

replace_dict = {
    "“": '"',
    "”": '"',
    "‘": "'",
    "’": "'",
    "—": "-",
    "…": "...",
    "《": "<",
    "》": ">",
    "（": "(",
    "）": ")",
    "【": "[",
    "】": "]",
    "、": ",",
    "。": ".",
    "，": ",",
    "：": ":",
    "；": ";",
    "？": "?",
    "！": "!",
    "\n": " "
}


# load dictionary
parser = CedictParser()
entries = parser.parse()
dictionary = {}
for e in entries:
    dictionary[e.simplified] = e.meanings[0]

c = Cihai()


class Character:

    def __init__(self, text, size):
        self.text = text
        self.size = size
        # check if character is chinese
        punctuation = re.findall('[{}]'.format(zhon.hanzi.punctuation), text)
        simplified = re.findall('[{}]'.format(zhon.cedict.simplified), text)
        self.chinese = len(punctuation + simplified) > 0
        self.color = RGBColor(0, 0, 0)  # black

    def len(self):
        len_ = len(self.text) * self.size
        if not self.chinese:
            return len_
        else:
            return len_ * ratio_lucida_to_mincho

    def set_color(self, color):
        self.color = color


def word_for_word(x, notranslate=False):
    if x in notranslate:
        return Character("", word_for_word_size)

    try:
        # ret = c.unihan.lookup_char(x).first().kDefinition.split(",")[0][:10]

        # dont translate some words

        # get first few full words that are shorter than 10 together
        splits = dictionary[x].split(" ")
        old_l = splits[0]
        for word in splits[1:]:

            new_l = old_l + " " + word
            if len(new_l) > 10:
                break

            old_l = new_l

        ret = old_l[:10]

        if not ret:
            ret = ""

    except:
        ret = ""

    return Character(ret, word_for_word_size)


def translation(x):
    # translate from chinese to english using a free translation service
    url = "https://translate.googleapis.com/translate_a/single?client=gtx&sl=zh-CN&tl=en&dt=t&q=" + x
    response = requests.get(url)
    if response.status_code != 200:
        raise Exception("ERROR: API request unsuccessful.")
    ret = json.loads(response.content.decode("utf-8"))[0][0][0]
    return Character(ret, translation_size)


def get_pinyin(x):
    # get pinyin from text x
    ret = pinyin.get(x)
    return Character(ret, pinyin_size)


def pad(x, length):
    if x.len() == length:
        return "", ""
    elif x.len() < length:
        # define left and right pad, round up left, round down right
        left = int(np.ceil((length - x.len()) / 2) / space_size)
        right = int(np.floor((length - x.len()) / 2) / space_size)

        return left * " ", right * " "


document = Document()
section = document.sections[0]
section.page_height = Mm(297 * 3)
section.page_width = Mm(210 * 3)


def preprocess(line):
    for key, value in replace_dict.items():
        line = line.replace(key, value)
    return line


def add_hyperlink(paragraph, text, url, color, font_name, size):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = color
    r.font.name = font_name
    r.font.size = size

    return hyperlink


def enrich_txt(path, use_notranslate_file=False):

    # load text file from data
    filename = "xwz.txt"
    with open(path, "r", encoding="UTF-8") as f:
        data = f.read()
        data = data.replace("。”", "”。")
        data = data.split("。")


    # load list of words that should not be translated
    if use_notranslate_file:
        with open("notranslate.txt", "r", encoding="UTF-8") as f:
            notranslate = f.read().splitlines()
            notranslate = set(notranslate)



    if not c.unihan.is_bootstrapped:  # download and install Unihan to db
        c.unihan.bootstrap()

    def display_padded_text(text_list, padding, add_dict_link=False):
        for i in range(len(text_list)):

            pad_left, pad_right = padding[i]
            run = p.add_run(pad_left)
            run.font.size = Pt(space_size)
            run.font.name = font_name
            # add hyperlink
            if add_dict_link:
                url = "https://www.mdbg.net/chinese/dictionary?page=worddict&wdrst=0&wdqb=" + text_list[i].text
                add_hyperlink(p, text_list[i].text, url, text_list[i].color, font_name, Pt(text_list[i].size))
            else:
                run = p.add_run(text_list[i].text)
                run.font.size = Pt(text_list[i].size)
                run.font.name = font_name
                run.font.color.rgb = text_list[i].color

            run = p.add_run(pad_right)
            run.font.size = Pt(space_size)
            run.font.name = font_name

            if debug:
                run = p.add_run("|")
                run.font.size = Pt(15)
                run.font.name = font_name

    line_i = 1
    for line in tqdm(data):
        line = preprocess(line)
        p = document.add_paragraph("")
        run = p.add_run(translation(line).text + "\n")
        run.font.size = Pt(translation_size)
        run.font.name = "Arial"
        wfw_string, pad_wfw = [], []
        py_string, pad_py = [], []
        char_string, pad_char = [], []
        import pynlpir

        pynlpir.open()

        tokenized = pynlpir.segment(line)
        sizes = []
        for token, word_type in tokenized:

            wfw = word_for_word(token, notranslate)
            py = get_pinyin(token)
            token = Character(token, base_size)
            token.set_color(color_coding.get(word_type, RGBColor(0, 0, 0)))
            size = max(wfw.len(), py.len(), token.len())
            sizes.append(size)
            pad_wfw.append(pad(wfw, size))
            pad_py.append(pad(py, size))
            pad_char.append(pad(token, size))

            wfw_string.append(wfw)
            py_string.append(py)

            char_string.append(token)

        c_l, c_r = 0, 0
        chunks = []
        while c_l < len(char_string) -1:
            sizes = np.array(sizes)
            cumsum = sizes[c_r:].cumsum()
            # get last value where cumsum is less than 2500

            c_r = c_l + np.where(cumsum < 2500)[0][-1]
            chunk = wfw_string[c_l:c_r], pad_wfw[c_l:c_r], py_string[c_l:c_r], pad_py[c_l:c_r], char_string[c_l:c_r], pad_char[c_l:c_r]
            chunks.append(chunk)
            c_l = c_r

        for wfw_string, pad_wfw, py_string, pad_py, char_string, pad_char in chunks:

            display_padded_text(wfw_string, pad_wfw)
            p.add_run("\n")
            display_padded_text(py_string, pad_py)
            p.add_run("\n")
            display_padded_text(char_string, pad_char, add_dict_link=True)


            if line_i % lines_per_page == 0:
                run = p.add_run()
                run.add_break(WD_BREAK.PAGE)
            else:
                p.add_run("\n")
            line_i+=1

    document.save(f'output/{filename}.docx')


if __name__ == "__main__":
    enrich_txt("data/xwz.txt")